"""Unit tests for the document parser.

Requirement: R9 — UI accepts PDF/DOCX/TXT uploads piped to Intake.
Module under test: app/tools/document_parser.py
"""
from __future__ import annotations

import pytest

from app.tools.document_parser import (
    DocumentTooLargeError,
    UnsupportedDocumentError,
    parse_document,
)


def test_parses_plain_text() -> None:
    out = parse_document("note.txt", b"Receipt: MYR 99.00")
    assert out.kind == "txt"
    assert "MYR 99.00" in out.text
    assert out.page_count == 0


def test_parses_markdown_as_text() -> None:
    out = parse_document("note.md", b"# Receipt\nTotal: MYR 12.50")
    assert out.kind == "txt"
    assert "Total: MYR 12.50" in out.text


def test_unsupported_extension_raises() -> None:
    with pytest.raises(UnsupportedDocumentError):
        parse_document("photo.jpg", b"\xff\xd8\xff\xe0")


def test_no_extension_raises() -> None:
    with pytest.raises(UnsupportedDocumentError):
        parse_document("data", b"hello")


def test_oversize_raises_with_max_bytes() -> None:
    with pytest.raises(DocumentTooLargeError):
        parse_document("big.txt", b"x" * 1000, max_bytes=100)


def test_to_dict_round_trips() -> None:
    out = parse_document("r.txt", b"abc")
    d = out.to_dict()
    assert d["filename"] == "r.txt"
    assert d["kind"] == "txt"
    assert d["text"] == "abc"
    assert d["bytes_read"] == 3


def test_text_decode_replaces_bad_bytes() -> None:
    # Invalid UTF-8 should not raise — replaced with the replacement char.
    out = parse_document("garbage.txt", b"\xff\xfe\xfdok")
    assert "ok" in out.text


# --- PDF and DOCX in-memory fixtures -----------------------------------------

def _minimal_pdf() -> bytes:
    import io
    from pypdf import PdfWriter
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=100)
    buf = io.BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _minimal_docx() -> bytes:
    import io
    from docx import Document
    doc = Document()
    doc.add_paragraph("Receipt: MYR 88.00")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "Item"
    tbl.cell(0, 1).text = "MYR 88.00"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_parses_valid_pdf() -> None:
    out = parse_document("receipt.pdf", _minimal_pdf())
    assert out.kind == "pdf"
    assert out.filename == "receipt.pdf"
    assert out.page_count == 1
    assert isinstance(out.text, str)
    assert out.bytes_read > 0


def test_parses_valid_docx_with_table() -> None:
    out = parse_document("receipt.docx", _minimal_docx())
    assert out.kind == "docx"
    assert out.page_count == 0
    assert "MYR 88.00" in out.text
    assert out.bytes_read > 0
