"""Extract plain text from uploaded documents (PDF, DOCX, TXT, MD).

Kept minimal on purpose: we only need to produce a receipt-text blob
that the Intake agent can read. OCR for scanned images is out of scope
for MVP — we accept digital PDFs only.
"""
from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path

from langsmith import traceable


SUPPORTED_EXTS = {".pdf", ".docx", ".txt", ".md"}


class UnsupportedDocumentError(ValueError):
    pass


class DocumentTooLargeError(ValueError):
    pass


@dataclass
class ParsedDocument:
    filename: str
    kind: str            # "pdf" | "docx" | "txt"
    text: str
    page_count: int      # 0 for non-paginated formats
    bytes_read: int

    def to_dict(self) -> dict:
        return {
            "filename": self.filename,
            "kind": self.kind,
            "text": self.text,
            "page_count": self.page_count,
            "bytes_read": self.bytes_read,
        }


def _parse_pdf(data: bytes, filename: str) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [p.extract_text() or "" for p in reader.pages]
    text = "\n\n".join(pages).strip()
    return ParsedDocument(
        filename=filename, kind="pdf",
        text=text, page_count=len(reader.pages), bytes_read=len(data),
    )


def _parse_docx(data: bytes, filename: str) -> ParsedDocument:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    # Also pull table cells — receipts often live in tables.
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text.strip())
    text = "\n".join(paragraphs).strip()
    return ParsedDocument(
        filename=filename, kind="docx",
        text=text, page_count=0, bytes_read=len(data),
    )


def _parse_text(data: bytes, filename: str) -> ParsedDocument:
    text = data.decode("utf-8", errors="replace").strip()
    return ParsedDocument(
        filename=filename, kind="txt",
        text=text, page_count=0, bytes_read=len(data),
    )


@traceable(run_type="tool", name="tool.parse_document")
def parse_document(filename: str, data: bytes, *, max_bytes: int | None = None) -> ParsedDocument:
    """Dispatch on file extension. Raises on unsupported/oversized input."""
    if max_bytes is not None and len(data) > max_bytes:
        raise DocumentTooLargeError(f"{filename} is {len(data)} bytes, max {max_bytes}")

    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(data, filename)
    if ext == ".docx":
        return _parse_docx(data, filename)
    if ext in (".txt", ".md"):
        return _parse_text(data, filename)
    raise UnsupportedDocumentError(
        f"Unsupported extension {ext!r}. Supported: {sorted(SUPPORTED_EXTS)}"
    )
